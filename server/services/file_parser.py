import fitz  # PyMuPDF
import docx
import io


def parse_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF bytes."""
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    text = ""
    for page in doc:
        text += page.get_text()
    doc.close()
    return text.strip()


def parse_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX bytes."""
    doc = docx.Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Also grab text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())
    return "\n".join(paragraphs)


def parse_resume_file(filename: str, file_bytes: bytes) -> str:
    """Route to the correct parser based on file extension."""
    ext = filename.rsplit(".", 1)[-1].lower()
    if ext == "pdf":
        return parse_pdf(file_bytes)
    elif ext == "docx":
        return parse_docx(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: {ext}")
